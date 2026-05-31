"""
md_to_docx.py — 將 final_report.md 轉換為格式化的 Word 文件

Usage:
    uv run final_project/md_to_docx.py
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH  = Path(__file__).parent / "output" / "final_report.md"
OUT_PATH = Path(__file__).parent / "output" / "final_report.docx"

CJK_FONT  = "微軟正黑體"
BODY_FONT = "微軟正黑體"
CODE_FONT = "Courier New"


# ── 字型輔助 ──────────────────────────────────────────────────────────

def set_run_font(run, font_name, size_pt=None, bold=False, italic=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt:
        run.font.size = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


# ── Token-based inline 解析器 ─────────────────────────────────────────
# 解決 ***bold**rest* / *整行斜體* / **粗體** 混合問題

def tokenize(text: str) -> list:
    """把 inline markdown 切成 tokens：('text', str) | ('marker', str) | ('code', str)"""
    tokens = []
    i = 0
    while i < len(text):
        ch = text[i]

        if ch == '`':
            j = text.find('`', i + 1)
            if j != -1:
                tokens.append(('code', text[i + 1:j]))
                i = j + 1
            else:
                tokens.append(('text', ch))
                i += 1

        elif ch == '*':
            # 先偵測連續 * 的長度
            j = i
            while j < len(text) and text[j] == '*':
                j += 1
            count = j - i
            # 最多認 3 個，多餘的當文字
            marker_len = min(count, 3)
            tokens.append(('marker', '*' * marker_len))
            if count > 3:
                tokens.append(('text', '*' * (count - 3)))
            i = j

        else:
            # 累積普通文字直到遇到 * 或 `
            j = i
            while j < len(text) and text[j] not in ('*', '`'):
                j += 1
            tokens.append(('text', text[i:j]))
            i = j

    return tokens


def add_inline(para, text: str, base_size=11, base_bold=False, base_italic=False):
    """解析 inline markdown 並加入 runs（支援嵌套 bold/italic/code）"""
    tokens = tokenize(text)

    bold   = base_bold
    italic = base_italic

    for tok_type, tok_val in tokens:
        if tok_type == 'text':
            if not tok_val:
                continue
            run = para.add_run(tok_val)
            set_run_font(run, BODY_FONT, base_size, bold=bold, italic=italic)

        elif tok_type == 'code':
            run = para.add_run(tok_val)
            # 使用等寬字型，紅色
            run.font.name = CODE_FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), CODE_FONT)
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = RGBColor(180, 0, 0)

        elif tok_type == 'marker':
            # Toggle: *** = bold+italic, ** = bold, * = italic
            if tok_val == '***':
                bold   = not bold
                italic = not italic
            elif tok_val == '**':
                bold   = not bold
            elif tok_val == '*':
                italic = not italic


# ── 段落建立輔助 ──────────────────────────────────────────────────────

def add_heading(doc, text: str, level: int):
    configs = {
        1: (18, True,  (31,  73,  125), WD_ALIGN_PARAGRAPH.CENTER),
        2: (14, True,  (31,  73,  125), WD_ALIGN_PARAGRAPH.LEFT),
        3: (12, True,  (68,  114, 196), WD_ALIGN_PARAGRAPH.LEFT),
        4: (11, True,  (0,   0,   0),   WD_ALIGN_PARAGRAPH.LEFT),
    }
    size, bold, color, align = configs.get(level, (11, True, (0, 0, 0), WD_ALIGN_PARAGRAPH.LEFT))
    clean = re.sub(r'^#+\s*', '', text).strip()
    para  = doc.add_heading('', level=min(level, 4))
    para.clear()
    run = para.add_run(clean)
    set_run_font(run, CJK_FONT, size, bold=bold, color=color)
    para.alignment = align
    para.paragraph_format.space_before = Pt(12 if level <= 2 else 6)
    para.paragraph_format.space_after  = Pt(4)
    return para


def add_body(doc, text: str, indent_cm=0, size=11):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.line_spacing = Pt(18)
    if indent_cm:
        para.paragraph_format.left_indent = Cm(indent_cm)
    add_inline(para, text.strip(), base_size=size)
    return para


def add_blockquote(doc, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Cm(1.2)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    # 左側灰線
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "12")
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), "AAAAAA")
    pBdr.append(left)
    pPr.append(pBdr)
    add_inline(para, text.strip(), base_size=10, base_italic=True)
    for run in para.runs:
        run.font.color.rgb = RGBColor(80, 80, 80)
    return para


def add_code_block(doc, lines: list):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Cm(0.8)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    run = para.add_run("\n".join(lines))
    run.font.name = CODE_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CODE_FONT)
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(60, 60, 60)
    return para


def add_table(doc, header_row: list, data_rows: list):
    n_cols = len(header_row)
    table  = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"

    # 表頭
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header_row):
        hdr_cells[i].text = ""
        _set_cell_bg(hdr_cells[i], "2E5496")
        p   = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 表頭用 add_inline，支援 **粗體**
        add_inline(p, h.strip(), base_size=10, base_bold=True)
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    # 資料列
    for ri, row in enumerate(data_rows):
        cells = table.add_row().cells
        bg    = "EAF0FB" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row):
            cells[ci].text = ""
            _set_cell_bg(cells[ci], bg)
            p = cells[ci].paragraphs[0]
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            add_inline(p, cell_text.strip(), base_size=10)

    doc.add_paragraph()
    return table


# ── Markdown 表格解析 ──────────────────────────────────────────────────

def parse_table_row(line: str) -> list:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator(cols: list) -> bool:
    return all(re.match(r'^[-: ]+$', c) for c in cols if c)


# ── 主轉換流程 ────────────────────────────────────────────────────────

def add_hrule(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "AAAAAA")
    pBdr.append(bot)
    pPr.append(pBdr)


def convert(md_path: Path, out_path: Path):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    lines     = md_path.read_text(encoding="utf-8").splitlines()
    in_code   = False
    code_buf  = []
    table_buf = []

    def flush_table():
        if not table_buf:
            return
        header = parse_table_row(table_buf[0])
        data   = []
        for row_line in table_buf[2:]:
            cols = parse_table_row(row_line)
            while len(cols) < len(header):
                cols.append("")
            data.append(cols[:len(header)])
        add_table(doc, header, data)
        table_buf.clear()

    i = 0
    while i < len(lines):
        line = lines[i]

        # ── 程式碼區塊 ──────────────────────────────────────────────
        if line.startswith("```"):
            if not in_code:
                flush_table()
                in_code  = True
                code_buf = []
            else:
                in_code = False
                add_code_block(doc, code_buf)
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── 表格列 ─────────────────────────────────────────────────
        if line.startswith("|"):
            table_buf.append(line)
            i += 1
            continue
        else:
            flush_table()

        # ── 空行 ────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── 水平線 ──────────────────────────────────────────────────
        if re.match(r'^-{3,}$', line.strip()):
            add_hrule(doc)
            i += 1
            continue

        # ── 標題 ────────────────────────────────────────────────────
        hm = re.match(r'^(#{1,4})\s+(.*)', line)
        if hm:
            add_heading(doc, hm.group(2), len(hm.group(1)))
            i += 1
            continue

        # ── 引用區塊 ────────────────────────────────────────────────
        if line.startswith(">"):
            add_blockquote(doc, re.sub(r'^>\s*', '', line))
            i += 1
            continue

        # ── LaTeX block（$$ ... $$） ──────────────────────────────
        if line.strip() == "$$":
            math_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != "$$":
                math_lines.append(lines[i].strip())
                i += 1
            i += 1
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run("  ".join(math_lines))
            run.font.name   = CODE_FONT
            run.font.size   = Pt(11)
            run.font.italic = True
            continue

        # ── 無序列表 ─────────────────────────────────────────────────
        bm = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if bm:
            indent = len(bm.group(1)) // 2 + 1
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent  = Cm(indent * 0.6)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            para.paragraph_format.line_spacing = Pt(16)
            add_inline(para, bm.group(2))
            i += 1
            continue

        # ── 有序列表 ─────────────────────────────────────────────────
        nm = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if nm:
            indent = len(nm.group(1)) // 2 + 1
            para = doc.add_paragraph(style="List Number")
            para.paragraph_format.left_indent  = Cm(indent * 0.6)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            para.paragraph_format.line_spacing = Pt(16)
            add_inline(para, nm.group(2))
            i += 1
            continue

        # ── 一般段落 ─────────────────────────────────────────────────
        add_body(doc, line)
        i += 1

    flush_table()

    doc.save(out_path)
    print(f"Done -> {out_path}")


if __name__ == "__main__":
    convert(MD_PATH, OUT_PATH)
